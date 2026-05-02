"""Tests for the document processing layer."""

from datetime import UTC, datetime
from pathlib import Path
from subprocess import CompletedProcess
from unittest.mock import patch

import pytest
from docx import Document

from src.documents.docx_engine import (
    build_resume,
    build_resume_from_ir,
    create_default_template,
    substitute_placeholders,
)
from src.documents.file_manager import get_output_paths, make_filename
from src.documents.latex_engine import (
    build_resume_tex_from_ir,
    compile_latex_to_pdf,
    latex_escape,
)
from src.documents.templates import (
    create_latex_template_package,
    discover_templates,
    ensure_template_package,
    get_template_package_detail,
    get_template_path,
    list_template_packages,
    register_template,
    save_uploaded_template_package,
    update_latex_template_package,
)
from src.generation.ir import ResumeBullet, ResumeDocument, ResumeItem

TMP_DIR = Path("data/output/_test")

SAMPLE_IDENTITY = {
    "full_name": "Jane Doe",
    "email": "jane@example.com",
    "phone": "+1 604-000-0000",
    "location": "Vancouver, BC",
    "linkedin_url": "linkedin.com/in/janedoe",
    "github_url": "github.com/janedoe",
}

SAMPLE_EDUCATION = [
    {
        "institution": "UBC",
        "degree": "Master of Science",
        "field": "Computer Science",
        "start_date": "2024-09",
        "end_date": "2026-04",
        "gpa": "4.0/4.0",
        "relevant_courses": [
            {"name": "Distributed Systems", "tags": ["distributed"]},
        ],
    }
]

SAMPLE_EXPERIENCES = [
    {
        "company": "Stripe",
        "title": "Software Engineer Intern",
        "location": "San Francisco, CA",
        "start_date": "2025-05",
        "end_date": "2025-08",
        "bullets": [
            {
                "text": "Built payment retry logic handling 1M+ transactions/day",
                "tags": ["backend"],
            },
        ],
    }
]

SAMPLE_PROJECTS = [
    {
        "name": "AutoApply",
        "role": "Lead Developer",
        "tech_stack": ["Python", "Playwright"],
        "bullets": [
            {
                "text": "Developed an AI agent for automated job applications",
                "tags": ["python", "ai"],
            },
        ],
    }
]

SAMPLE_SKILLS = {
    "languages": ["Python", "Java"],
    "frameworks": ["FastAPI", "React"],
    "databases": ["PostgreSQL"],
    "tools": ["Docker"],
}


@pytest.fixture(autouse=True)
def cleanup():
    yield
    import shutil

    if TMP_DIR.exists():
        shutil.rmtree(TMP_DIR)


class TestDocxEngine:
    def test_create_default_template(self):
        template_path = TMP_DIR / "default_template.docx"
        result = create_default_template(template_path)
        assert result.exists()
        doc = Document(str(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "{{FULL_NAME}}" in full_text
        assert "{{EDUCATION_BLOCK}}" in full_text
        assert "{{EXPERIENCE_BLOCK}}" in full_text

    def test_substitute_placeholders(self):
        template_path = TMP_DIR / "template.docx"
        create_default_template(template_path)
        doc = Document(str(template_path))
        substitute_placeholders(doc, {"FULL_NAME": "Jane Doe", "EMAIL": "jane@example.com"})
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text
        assert "jane@example.com" in full_text

    def test_build_resume(self):
        template_path = TMP_DIR / "template.docx"
        create_default_template(template_path)
        output_path = TMP_DIR / "resume_test.docx"

        result = build_resume(
            template_path=template_path,
            identity=SAMPLE_IDENTITY,
            education=SAMPLE_EDUCATION,
            experiences=SAMPLE_EXPERIENCES,
            projects=SAMPLE_PROJECTS,
            skills=SAMPLE_SKILLS,
            selected_bullets={},
            output_path=output_path,
        )

        assert result.exists()
        doc = Document(str(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text
        assert "Stripe" in full_text
        assert "AutoApply" in full_text
        assert "Python" in full_text
        assert "UBC" in full_text

    def test_build_resume_from_ir(self):
        template_path = TMP_DIR / "template_ir.docx"
        create_default_template(template_path)
        output_path = TMP_DIR / "resume_ir.docx"
        document = ResumeDocument(
            target_role="Backend Intern",
            company="Stripe",
            header=SAMPLE_IDENTITY,
            education=SAMPLE_EDUCATION,
            skills=SAMPLE_SKILLS,
            section_order=["header", "projects", "skills", "experience", "education"],
            experiences=[
                ResumeItem(
                    source_id="experience:stripe",
                    source_type="experience",
                    name="Stripe",
                    organization="Stripe",
                    title="Software Engineer Intern",
                    location="San Francisco, CA",
                    start_date="2025-05",
                    end_date="2025-08",
                    bullets=[
                        ResumeBullet(
                            text="Built payment retry logic handling 1M+ transactions/day",
                            source_id="experience:stripe:bullet:0",
                            source_type="experience",
                            source_entity="Stripe - Software Engineer Intern",
                        )
                    ],
                )
            ],
            projects=[],
        )

        result = build_resume_from_ir(template_path, document, output_path)

        assert result.exists()
        doc = Document(str(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in full_text
        assert "Stripe" in full_text
        assert "payment retry" in full_text
        paragraph_text = [p.text for p in doc.paragraphs]
        assert paragraph_text.index("Skills") < paragraph_text.index("Experience")
        assert paragraph_text.index("Experience") < paragraph_text.index("Education")

    def test_template_package_renderer_uses_named_styles(self, tmp_path):
        package = ensure_template_package("resume", template_root=tmp_path)
        output_path = tmp_path / "resume_named_styles.docx"
        document = ResumeDocument(
            target_role="Backend Intern",
            company="Stripe",
            header=SAMPLE_IDENTITY,
            education=SAMPLE_EDUCATION,
            skills=SAMPLE_SKILLS,
            section_order=["header", "skills"],
            experiences=[],
            projects=[],
        )

        result = build_resume_from_ir(
            package.template_path,
            document,
            output_path,
            manifest=package.manifest,
        )

        doc = Document(str(result))
        styles_by_text = {paragraph.text: paragraph.style.name for paragraph in doc.paragraphs}
        assert styles_by_text["Jane Doe"] == "Resume.Name"
        assert styles_by_text["Skills"] == "Resume.SectionHeading"
        assert "{{resume.sections}}" not in " ".join(p.text for p in doc.paragraphs)


class TestFileManager:
    def test_make_filename_resume(self):
        date = datetime(2026, 4, 2, tzinfo=UTC)
        name = make_filename("resume", "Stripe", "Backend Engineer", date)
        assert name == "resume_stripe_backend_engineer_2026-04-02.docx"

    def test_make_filename_cover(self):
        date = datetime(2026, 4, 2, tzinfo=UTC)
        name = make_filename("cover", "Google LLC", "SWE Intern", date, ext="pdf")
        assert name == "cover_google_llc_swe_intern_2026-04-02.pdf"

    def test_make_filename_special_chars(self):
        date = datetime(2026, 4, 2, tzinfo=UTC)
        name = make_filename("resume", "A/B & Co.", "C++ Dev", date)
        # Should not contain special chars
        assert "/" not in name
        assert "&" not in name

    def test_get_output_paths(self):
        date = datetime(2026, 4, 2, tzinfo=UTC)
        paths = get_output_paths(TMP_DIR, "Stripe", "Backend Intern", date)
        assert "resume_docx" in paths
        assert "resume_pdf" in paths
        assert "resume_tex" in paths
        assert "cover_docx" in paths
        assert "cover_pdf" in paths
        assert "cover_tex" in paths
        assert paths["resume_docx"].suffix == ".docx"
        assert paths["resume_pdf"].suffix == ".pdf"
        assert paths["resume_tex"].suffix == ".tex"


class TestTemplateRegistry:
    def test_register_and_get(self):
        template_path = TMP_DIR / "my_template.docx"
        create_default_template(template_path)
        register_template("test_template", template_path)
        assert get_template_path("test_template") == template_path

    def test_get_missing_template_raises(self):
        with pytest.raises(KeyError):
            get_template_path("nonexistent_template_xyz")

    def test_discover_templates(self):
        # Create two templates
        create_default_template(TMP_DIR / "modern.docx")
        create_default_template(TMP_DIR / "classic.docx")
        discover_templates(TMP_DIR)
        # Both should now be registered
        assert get_template_path("modern") is not None
        assert get_template_path("classic") is not None

    def test_list_template_packages(self, tmp_path):
        resume_package = ensure_template_package("resume", template_root=tmp_path)
        ensure_template_package("cover_letter", template_root=tmp_path)
        (resume_package.directory / "preview.pdf").write_bytes(b"pdf")

        templates = list_template_packages(template_root=tmp_path)

        assert templates["resume"][0]["template_id"] == "ats_single_column_v1"
        assert templates["resume"][0]["preview_pdf"] == "preview.pdf"
        assert not Path(templates["resume"][0]["preview_pdf"]).is_absolute()
        assert templates["resume"][0]["validation"]["ok"] is True
        assert templates["cover_letter"][0]["template_id"] == "classic_v1"

    def test_save_uploaded_template_package(self, tmp_path):
        upload_docx = tmp_path / "upload.docx"
        doc = Document()
        doc.add_paragraph("Uploaded resume template")
        doc.save(str(upload_docx))

        template = save_uploaded_template_package(
            document_type="resume",
            filename="upload.docx",
            content=upload_docx.read_bytes(),
            template_name="My Resume Template",
            template_root=tmp_path / "templates",
        )

        assert template["template_id"] == "my_resume_template"
        assert template["validation"]["ok"] is True

    def test_template_package_rejects_path_traversal_id(self, tmp_path):
        with pytest.raises(ValueError, match="Invalid template id"):
            ensure_template_package(
                "resume",
                "..\\escape",
                template_root=tmp_path / "templates",
            )

        assert not (tmp_path / "escape").exists()


class TestLatexTemplates:
    def test_create_latex_template_package(self, tmp_path):
        template = create_latex_template_package(
            document_type="resume",
            template_name="Technical LaTeX Resume",
            template_root=tmp_path / "templates",
        )

        assert template["template_id"] == "technical_latex_resume"
        assert template["renderer"] == "latex"
        assert template["supported_outputs"] == ["tex", "pdf"]
        assert template["validation"]["ok"] is True

        detail = get_template_package_detail(
            "resume",
            template["template_id"],
            template_root=tmp_path / "templates",
        )
        assert "{{resume.sections}}" in detail["content"]

    def test_uploaded_latex_template_reports_missing_marker(self, tmp_path):
        template = save_uploaded_template_package(
            document_type="resume",
            filename="plain.tex",
            content=b"\\documentclass{article}\n\\begin{document}No marker\\end{document}\n",
            template_name="Plain LaTeX",
            template_root=tmp_path / "templates",
        )

        assert template["renderer"] == "latex"
        assert template["validation"]["ok"] is False
        assert template["validation"]["issues"][0]["type"] == "missing_block"
        assert template["validation"]["issues"][0]["severity"] == "error"
        assert "Add this marker exactly" in template["validation"]["issues"][0]["message"]

    def test_update_latex_template_package(self, tmp_path):
        template = create_latex_template_package(
            document_type="cover_letter",
            template_name="Editable Cover",
            template_root=tmp_path / "templates",
        )
        updated = update_latex_template_package(
            document_type="cover_letter",
            template_id=template["template_id"],
            template_name="Updated Cover",
            description="Edited in tests.",
            content="\\documentclass{article}\n\\begin{document}\n{{cover_letter.body}}\n\\end{document}\n",
            template_root=tmp_path / "templates",
        )

        assert updated["name"] == "Updated Cover"
        assert updated["description"] == "Edited in tests."
        assert updated["validation"]["ok"] is True

    def test_latex_escape(self):
        escaped = latex_escape(r"R&D_50% C# {x} \ path")

        assert r"R\&D\_50\% C\# \{x\}" in escaped
        assert r"\textbackslash{} path" in escaped

    def test_build_resume_tex_from_ir(self, tmp_path):
        template = create_latex_template_package(
            document_type="resume",
            template_name="Render LaTeX Resume",
            template_root=tmp_path / "templates",
        )
        package = ensure_template_package(
            "resume",
            template["template_id"],
            template_root=tmp_path / "templates",
        )
        output_path = tmp_path / "resume.tex"
        document = ResumeDocument(
            target_role="Backend Intern",
            company="Stripe",
            header={**SAMPLE_IDENTITY, "full_name": "Jane & Doe"},
            education=SAMPLE_EDUCATION,
            skills=SAMPLE_SKILLS,
            section_order=["header", "skills", "experience", "education"],
            experiences=[
                ResumeItem(
                    source_id="experience:stripe",
                    source_type="experience",
                    name="Stripe",
                    organization="Stripe",
                    title="Software Engineer Intern",
                    bullets=[
                        ResumeBullet(
                            text="Built R&D tooling with C# and 50% less toil",
                            source_id="experience:stripe:bullet:0",
                            source_type="experience",
                            source_entity="Stripe",
                        )
                    ],
                )
            ],
            projects=[],
        )

        result = build_resume_tex_from_ir(
            package.template_path,
            document,
            output_path,
            manifest=package.manifest,
        )

        text = result.read_text(encoding="utf-8")
        assert r"Jane \& Doe" in text
        assert r"\section*{Skills}" in text
        assert r"R\&D tooling with C\# and 50\% less toil" in text
        assert "{{resume.sections}}" not in text

    def test_latex_renderer_rejects_missing_marker(self, tmp_path):
        template = create_latex_template_package(
            document_type="resume",
            template_name="Broken LaTeX Resume",
            template_root=tmp_path / "templates",
        )
        package = ensure_template_package(
            "resume",
            template["template_id"],
            template_root=tmp_path / "templates",
        )
        package.template_path.write_text(
            "\\documentclass{article}\n\\begin{document}No marker\\end{document}\n",
            encoding="utf-8",
        )

        with pytest.raises(ValueError, match="missing block marker"):
            build_resume_tex_from_ir(
                package.template_path,
                ResumeDocument(
                    target_role="Backend Intern",
                    company="Stripe",
                    header=SAMPLE_IDENTITY,
                ),
                tmp_path / "out.tex",
                manifest=package.manifest,
            )

    def test_compile_latex_to_pdf_disables_shell_escape(self, tmp_path):
        tex_path = tmp_path / "resume.tex"
        tex_path.write_text("\\documentclass{article}\n\\begin{document}Hi\\end{document}\n")
        pdf_path = tmp_path / "resume.pdf"

        def fake_run(command, *, cwd, **kwargs):
            Path(cwd, "main.pdf").write_bytes(b"%PDF")
            Path(cwd, "main.log").write_text("compile log", encoding="utf-8")
            return CompletedProcess(command, 0, stdout="", stderr="")

        def fake_which(name):
            return "pdflatex" if name == "pdflatex" else None

        with (
            patch("src.documents.latex_engine.shutil.which", side_effect=fake_which),
            patch("src.documents.latex_engine.subprocess.run", side_effect=fake_run) as run,
        ):
            result = compile_latex_to_pdf(tex_path, pdf_path)

        assert result == pdf_path
        assert pdf_path.exists()
        assert pdf_path.with_suffix(".log").read_text(encoding="utf-8") == "compile log"
        first_command = run.call_args_list[0].args[0]
        assert "-no-shell-escape" in first_command
