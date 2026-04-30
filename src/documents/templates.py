"""Document template package management.

Templates are first-class assets: a DOCX file owns visual style definitions,
while manifest.json describes named styles and content capacity constraints.
"""

from __future__ import annotations

import json
import logging
import re
import shutil
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
from pydantic import BaseModel, Field

from src.core.config import PROJECT_ROOT

logger = logging.getLogger("autoapply.documents.templates")

TEMPLATE_REGISTRY: dict[str, Path] = {}
TEMPLATE_ROOT = PROJECT_ROOT / "data" / "templates"
DEFAULT_TEMPLATE_IDS = {
    "resume": "ats_single_column_v1",
    "cover_letter": "classic_v1",
}
_TEMPLATE_ID_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9_-]{0,99}$")

DocumentType = Literal["resume", "cover_letter"]
TemplateFormat = Literal["docx", "latex"]
TemplateRenderer = Literal["docx", "latex"]
TemplateOutput = Literal["docx", "pdf", "tex"]


class TemplatePage(BaseModel):
    size: str = "letter"
    max_pages: int = 1


class TemplateSection(BaseModel):
    enabled: bool = True
    max_items: int | None = None
    max_bullets_per_item: int | None = None
    max_words_per_bullet: int | None = None
    max_lines: int | None = None


class TemplateCapacity(BaseModel):
    max_pages: int = 1
    max_sections: int | None = None
    max_experience_items: int | None = None
    max_project_items: int | None = None
    max_bullets_total: int | None = None
    max_words_per_bullet: int | None = None
    max_skill_lines: int | None = None


class TemplateManifest(BaseModel):
    template_id: str
    document_type: DocumentType
    template_format: TemplateFormat = "docx"
    renderer: TemplateRenderer = "docx"
    supported_outputs: list[TemplateOutput] = Field(default_factory=lambda: ["docx", "pdf"])
    name: str = ""
    description: str = ""
    page: TemplatePage = Field(default_factory=TemplatePage)
    styles: dict[str, str] = Field(default_factory=dict)
    sections: dict[str, TemplateSection] = Field(default_factory=dict)
    section_order: list[str] = Field(default_factory=list)
    capacity: TemplateCapacity = Field(default_factory=TemplateCapacity)
    blocks: dict[str, str] = Field(default_factory=dict)


class TemplatePackage(BaseModel):
    template_id: str
    document_type: DocumentType
    directory: Path
    template_path: Path
    manifest_path: Path
    manifest: TemplateManifest


def register_template(name: str, path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(f"Template not found: {path}")
    TEMPLATE_REGISTRY[name] = path
    logger.debug("Registered template '%s' at %s", name, path)


def get_template_path(name: str) -> Path:
    if name not in TEMPLATE_REGISTRY:
        raise KeyError(f"Template '{name}' not registered. Available: {list(TEMPLATE_REGISTRY)}")
    return TEMPLATE_REGISTRY[name]


def discover_templates(template_dir: Path) -> None:
    """Auto-register all .docx files found in template_dir."""
    if not template_dir.exists():
        logger.warning("Template directory not found: %s", template_dir)
        return
    for path in template_dir.glob("*.docx"):
        register_template(path.stem, path)
    logger.info("Discovered %d templates in %s", len(TEMPLATE_REGISTRY), template_dir)


def ensure_template_package(
    document_type: DocumentType,
    template_id: str | None = None,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> TemplatePackage:
    """Create the default template package if needed and return it."""
    template_id = template_id or DEFAULT_TEMPLATE_IDS[document_type]
    package_dir = _template_package_dir(document_type, template_id, template_root)
    package_dir.mkdir(parents=True, exist_ok=True)

    manifest_path = package_dir / "manifest.json"

    if manifest_path.exists():
        package = load_template_package(document_type, template_id, template_root=template_root)
        _ensure_required_markers(package)
        _write_sample_assets(package)
        return load_template_package(document_type, template_id, template_root=template_root)

    template_path = package_dir / "template.docx"

    manifest_path.write_text(
        json.dumps(_default_manifest(document_type, template_id), indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    if not template_path.exists():
        if document_type == "resume":
            _create_default_resume_template(template_path)
        else:
            _create_default_cover_letter_template(template_path)

    package = load_template_package(document_type, template_id, template_root=template_root)
    _ensure_required_markers(package)
    _write_sample_assets(package)
    return load_template_package(document_type, template_id, template_root=template_root)


def list_template_packages(
    document_type: DocumentType | None = None,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> dict[str, list[dict]]:
    """Return discovered template packages grouped by document type."""
    document_types = [document_type] if document_type else ["resume", "cover_letter"]
    grouped: dict[str, list[dict]] = {kind: [] for kind in document_types}
    for kind in document_types:
        ensure_template_package(kind, template_root=template_root)
        kind_dir = template_root / kind
        for manifest_path in sorted(kind_dir.glob("*/manifest.json")):
            template_id = manifest_path.parent.name
            try:
                package = load_template_package(kind, template_id, template_root=template_root)
            except Exception as exc:
                logger.warning(
                    "Skipping invalid template package %s: %s",
                    manifest_path.parent,
                    exc,
                )
                continue
            grouped[kind].append(serialize_template_package(package))
    return grouped


def save_uploaded_template_package(
    *,
    document_type: DocumentType,
    filename: str,
    content: bytes,
    template_name: str | None = None,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Persist an uploaded DOCX or single-file LaTeX template package."""
    if document_type not in DEFAULT_TEMPLATE_IDS:
        raise ValueError("Unsupported template document type.")
    suffix = Path(filename).suffix.lower()
    if suffix not in {".docx", ".tex"}:
        raise ValueError("Only .docx and .tex templates are supported.")
    if not content:
        raise ValueError("Template upload is empty.")

    if suffix == ".tex":
        return _save_uploaded_latex_template_package(
            document_type=document_type,
            filename=filename,
            content=content,
            template_name=template_name,
            template_root=template_root,
        )

    return _save_uploaded_docx_template_package(
        document_type=document_type,
        filename=filename,
        content=content,
        template_name=template_name,
        template_root=template_root,
    )


def create_latex_template_package(
    *,
    document_type: DocumentType,
    template_name: str | None = None,
    description: str | None = None,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Create a blank single-file LaTeX template package and return metadata."""
    if document_type not in DEFAULT_TEMPLATE_IDS:
        raise ValueError("Unsupported template document type.")

    default_name = "LaTeX Resume Template" if document_type == "resume" else "LaTeX Cover Letter"
    display_name = (template_name or default_name).strip() or default_name
    template_id = _unique_template_id(template_root / document_type, _slugify(display_name))
    package_dir = template_root / document_type / template_id
    package_dir.mkdir(parents=True, exist_ok=False)

    template_path = package_dir / "template.tex"
    manifest_path = package_dir / "manifest.json"
    try:
        manifest_payload = _default_latex_manifest(document_type, template_id)
        manifest_payload["name"] = display_name
        if description is not None:
            manifest_payload["description"] = description.strip()
        template_path.write_text(
            _default_latex_template(document_type),
            encoding="utf-8",
            newline="\n",
        )
        manifest_path.write_text(
            json.dumps(manifest_payload, indent=2) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        package = load_template_package(document_type, template_id, template_root=template_root)
        _write_sample_assets(package)
        return serialize_template_package(package)
    except Exception:
        shutil.rmtree(package_dir, ignore_errors=True)
        raise


def get_template_package_detail(
    document_type: DocumentType,
    template_id: str,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Return template metadata plus editable content for text-based templates."""
    package = load_template_package(document_type, template_id, template_root=template_root)
    serialized = serialize_template_package(package)
    serialized["content"] = (
        package.template_path.read_text(encoding="utf-8")
        if package.manifest.renderer == "latex"
        else None
    )
    return serialized


def update_latex_template_package(
    *,
    document_type: DocumentType,
    template_id: str,
    content: str,
    template_name: str | None = None,
    description: str | None = None,
    template_root: Path = TEMPLATE_ROOT,
) -> dict:
    """Update editable LaTeX template content and metadata."""
    package = load_template_package(document_type, template_id, template_root=template_root)
    if package.manifest.renderer != "latex":
        raise ValueError("Only LaTeX templates can be edited as text.")
    if not content.strip():
        raise ValueError("Template content is required.")

    package.template_path.write_text(content, encoding="utf-8", newline="\n")
    manifest = package.manifest.model_copy(
        update={
            "name": template_name.strip() if template_name is not None else package.manifest.name,
            "description": (
                description.strip() if description is not None else package.manifest.description
            ),
        }
    )
    package.manifest_path.write_text(
        json.dumps(manifest.model_dump(mode="json"), indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    package = load_template_package(document_type, template_id, template_root=template_root)
    _write_sample_assets(package)
    return serialize_template_package(package)


def _save_uploaded_docx_template_package(
    *,
    document_type: DocumentType,
    filename: str,
    content: bytes,
    template_name: str | None,
    template_root: Path,
) -> dict:
    """Persist an uploaded DOCX as a template package and return metadata."""

    display_name = (template_name or Path(filename).stem or "Uploaded Template").strip()
    template_id = _unique_template_id(template_root / document_type, _slugify(display_name))
    package_dir = template_root / document_type / template_id
    package_dir.mkdir(parents=True, exist_ok=False)

    template_path = package_dir / "template.docx"
    manifest_path = package_dir / "manifest.json"
    try:
        template_path.write_bytes(content)
        doc = Document(str(template_path))
        manifest_payload = _default_manifest(document_type, template_id)
        manifest_payload["name"] = display_name
        manifest_payload["description"] = f"Uploaded {document_type.replace('_', ' ')} template."
        manifest_path.write_text(
            json.dumps(manifest_payload, indent=2) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        manifest = TemplateManifest.model_validate(manifest_payload)
        _ensure_default_styles(doc, document_type)
        doc.save(str(template_path))
        package = TemplatePackage(
            template_id=template_id,
            document_type=document_type,
            directory=package_dir,
            template_path=template_path,
            manifest_path=manifest_path,
            manifest=manifest,
        )
        _ensure_required_markers(package)
        package = load_template_package(document_type, template_id, template_root=template_root)
        _write_sample_assets(package)
        return serialize_template_package(package)
    except Exception:
        shutil.rmtree(package_dir, ignore_errors=True)
        raise


def _save_uploaded_latex_template_package(
    *,
    document_type: DocumentType,
    filename: str,
    content: bytes,
    template_name: str | None,
    template_root: Path,
) -> dict:
    """Persist an uploaded single-file LaTeX template package and return metadata."""
    try:
        template_text = content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ValueError("LaTeX templates must be UTF-8 text files.") from exc
    if "\x00" in template_text:
        raise ValueError("LaTeX template contains invalid null bytes.")

    display_name = (template_name or Path(filename).stem or "Uploaded LaTeX Template").strip()
    template_id = _unique_template_id(template_root / document_type, _slugify(display_name))
    package_dir = template_root / document_type / template_id
    package_dir.mkdir(parents=True, exist_ok=False)

    template_path = package_dir / "template.tex"
    manifest_path = package_dir / "manifest.json"
    try:
        manifest_payload = _default_latex_manifest(document_type, template_id)
        manifest_payload["name"] = display_name
        manifest_payload["description"] = (
            f"Uploaded {document_type.replace('_', ' ')} LaTeX template."
        )
        template_path.write_text(template_text, encoding="utf-8", newline="\n")
        manifest_path.write_text(
            json.dumps(manifest_payload, indent=2) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        package = load_template_package(document_type, template_id, template_root=template_root)
        _write_sample_assets(package)
        return serialize_template_package(package)
    except Exception:
        shutil.rmtree(package_dir, ignore_errors=True)
        raise


def serialize_template_package(package: TemplatePackage) -> dict:
    """Serialize template metadata for APIs/UI without leaking system internals."""
    preview_pdf = package.directory / "preview.pdf"
    preview_png = package.directory / "preview.png"
    return {
        "template_id": package.template_id,
        "document_type": package.document_type,
        "template_format": package.manifest.template_format,
        "renderer": package.manifest.renderer,
        "supported_outputs": package.manifest.supported_outputs,
        "name": package.manifest.name or package.template_id.replace("_", " ").title(),
        "description": package.manifest.description,
        "manifest": package.manifest.model_dump(mode="json"),
        "preview_pdf": _public_asset_path(preview_pdf),
        "preview_png": _public_asset_path(preview_png),
        "validation": validate_template_package(package),
    }


def validate_template_package(package: TemplatePackage) -> dict:
    """Check the template has the styles and block markers declared by manifest."""
    issues = []
    try:
        if package.manifest.renderer == "latex":
            text = package.template_path.read_text(encoding="utf-8")
            if not text.strip():
                issues.append({"type": "empty_template", "message": "Template content is empty."})
        else:
            doc = Document(str(package.template_path))
            style_names = {style.name for style in doc.styles}
            for style in package.manifest.styles.values():
                if style not in style_names:
                    issues.append({"type": "missing_style", "message": f"Missing style: {style}"})
            text = _document_text(doc)
        for marker in package.manifest.blocks.values():
            if marker and marker not in text:
                issues.append(
                    {"type": "missing_block", "message": f"Missing block marker: {marker}"}
                )
    except Exception as exc:
        issues.append({"type": "template_unreadable", "message": str(exc)})
    return {"ok": not issues, "issues": issues}


def load_template_package(
    document_type: DocumentType,
    template_id: str | None = None,
    *,
    template_root: Path = TEMPLATE_ROOT,
) -> TemplatePackage:
    """Load a template package from templates/<document_type>/<template_id>."""
    template_id = template_id or DEFAULT_TEMPLATE_IDS[document_type]
    package_dir = _template_package_dir(document_type, template_id, template_root)
    manifest_path = package_dir / "manifest.json"

    if not manifest_path.exists():
        raise FileNotFoundError(f"Template manifest not found: {manifest_path}")

    manifest = TemplateManifest.model_validate_json(manifest_path.read_text(encoding="utf-8"))
    template_name = "template.tex" if manifest.renderer == "latex" else "template.docx"
    template_path = package_dir / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")
    if manifest.document_type != document_type:
        raise ValueError(
            f"Template {template_id} is {manifest.document_type}, not {document_type}"
        )
    return TemplatePackage(
        template_id=template_id,
        document_type=document_type,
        directory=package_dir,
        template_path=template_path,
        manifest_path=manifest_path,
        manifest=manifest,
    )


def default_manifest(document_type: Literal["resume", "cover_letter"]) -> TemplateManifest:
    template_id = DEFAULT_TEMPLATE_IDS[document_type]
    return TemplateManifest.model_validate(_default_manifest(document_type, template_id))


def _default_manifest(document_type: str, template_id: str) -> dict:
    if document_type == "cover_letter":
        return {
            "template_id": template_id,
            "document_type": "cover_letter",
            "template_format": "docx",
            "renderer": "docx",
            "supported_outputs": ["docx", "pdf"],
            "name": "Classic Cover Letter",
            "description": "Simple one-page cover letter with editable Word styles.",
            "page": {"size": "letter", "max_pages": 1},
            "styles": {
                "header": "CoverLetter.Header",
                "date": "CoverLetter.Date",
                "recipient": "CoverLetter.Recipient",
                "body": "CoverLetter.Body",
                "signature": "CoverLetter.Signature",
            },
            "sections": {
                "header": {"enabled": True},
                "body": {"enabled": True, "max_items": 5},
            },
            "section_order": ["header", "recipient", "body", "signature"],
            "capacity": {"max_pages": 1, "max_sections": 4},
            "blocks": {"body": "{{cover_letter.body}}"},
        }

    return {
        "template_id": template_id,
        "document_type": "resume",
        "template_format": "docx",
        "renderer": "docx",
        "supported_outputs": ["docx", "pdf"],
        "name": "ATS Single Column",
        "description": "One-page ATS-friendly resume with named Word styles and tab-stop dates.",
        "page": {"size": "letter", "max_pages": 1},
        "styles": {
            "name": "Resume.Name",
            "contact": "Resume.Contact",
            "section_heading": "Resume.SectionHeading",
            "item_title": "Resume.ItemTitle",
            "item_subtitle": "Resume.ItemSubtitle",
            "item_meta": "Resume.ItemMeta",
            "normal": "Resume.Normal",
            "bullet": "Resume.Bullet",
            "skill_category": "Resume.SkillCategory",
            "skill_line": "Resume.SkillLine",
        },
        "sections": {
            "summary": {"enabled": True, "max_items": 1},
            "education": {"enabled": True, "max_items": 2},
            "experience": {
                "enabled": True,
                "max_items": 3,
                "max_bullets_per_item": 4,
                "max_words_per_bullet": 24,
            },
            "projects": {
                "enabled": True,
                "max_items": 3,
                "max_bullets_per_item": 3,
                "max_words_per_bullet": 22,
            },
            "skills": {"enabled": True, "max_lines": 4},
        },
        "section_order": ["header", "education", "skills", "projects", "experience"],
        "capacity": {
            "max_pages": 1,
            "max_sections": 5,
            "max_experience_items": 3,
            "max_project_items": 3,
            "max_bullets_total": 13,
            "max_words_per_bullet": 24,
            "max_skill_lines": 4,
        },
        "blocks": {"sections": "{{resume.sections}}"},
    }


def _default_latex_manifest(document_type: str, template_id: str) -> dict:
    manifest = _default_manifest(document_type, template_id)
    manifest["template_format"] = "latex"
    manifest["renderer"] = "latex"
    manifest["supported_outputs"] = ["tex", "pdf"]
    manifest["styles"] = {}
    if document_type == "cover_letter":
        manifest["name"] = "LaTeX Cover Letter"
        manifest["description"] = "Single-file LaTeX cover letter template."
    else:
        manifest["name"] = "LaTeX Resume"
        manifest["description"] = "Single-file LaTeX resume template."
    return manifest


def _default_latex_template(document_type: str) -> str:
    if document_type == "cover_letter":
        return r"""\documentclass[11pt,letterpaper]{article}
\usepackage[margin=0.8in]{geometry}
\usepackage[hidelinks]{hyperref}
\setlength{\parindent}{0pt}
\setlength{\parskip}{8pt}

\begin{document}

{\large\textbf{ {{applicant.name}} }}\\
{{applicant.contact}}

{{recipient.company}}

{{cover_letter.body}}

Sincerely,\\
{{signature}}

\end{document}
"""

    return r"""\documentclass[10pt,letterpaper]{article}
\usepackage[margin=0.65in]{geometry}
\usepackage[hidelinks]{hyperref}
\setlength{\parindent}{0pt}
\setlength{\parskip}{4pt}

\begin{document}

{\LARGE\textbf{ {{full_name}} }}\\
{{contact}}\\
{{links}}

{{resume.sections}}

\end{document}
"""


def _create_default_resume_template(path: Path) -> None:
    doc = Document()
    _set_page(doc)
    _ensure_default_styles(doc, "resume")
    doc.add_paragraph("{{full_name}}", style="Resume.Name")
    doc.add_paragraph("{{contact}}", style="Resume.Contact")
    doc.add_paragraph("{{links}}", style="Resume.Contact")
    doc.add_paragraph("{{resume.sections}}", style="Resume.Normal")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _create_default_cover_letter_template(path: Path) -> None:
    doc = Document()
    _set_page(doc)
    _ensure_default_styles(doc, "cover_letter")
    doc.add_paragraph("{{applicant.name}}", style="CoverLetter.Header")
    doc.add_paragraph("{{applicant.contact}}", style="CoverLetter.Header")
    doc.add_paragraph("{{recipient.company}}", style="CoverLetter.Recipient")
    doc.add_paragraph("{{cover_letter.body}}", style="CoverLetter.Body")
    doc.add_paragraph("{{signature}}", style="CoverLetter.Signature")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _ensure_style(
    doc: Document,
    name: str,
    *,
    size: int,
    bold: bool = False,
    base: str | None = None,
    space_before: int = 0,
    space_after: int = 0,
    right_tab: bool = False,
) -> None:
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        try:
            style.base_style = doc.styles[base]
        except KeyError:
            pass
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    if right_tab:
        style.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)


def _ensure_default_styles(doc: Document, document_type: str) -> None:
    if document_type == "resume":
        _ensure_style(doc, "Resume.Name", size=16, bold=True)
        _ensure_style(doc, "Resume.Contact", size=9)
        _ensure_style(doc, "Resume.SectionHeading", size=11, bold=True, space_before=6)
        _ensure_style(doc, "Resume.ItemTitle", size=10, bold=True, right_tab=True)
        _ensure_style(doc, "Resume.ItemSubtitle", size=9)
        _ensure_style(doc, "Resume.ItemMeta", size=9)
        _ensure_style(doc, "Resume.Normal", size=9)
        _ensure_style(doc, "Resume.SkillCategory", size=9, bold=True)
        _ensure_style(doc, "Resume.SkillLine", size=9)
        _ensure_style(doc, "Resume.Bullet", size=9, base="List Bullet")
        return
    _ensure_style(doc, "CoverLetter.Header", size=10)
    _ensure_style(doc, "CoverLetter.Date", size=10)
    _ensure_style(doc, "CoverLetter.Recipient", size=10)
    _ensure_style(doc, "CoverLetter.Body", size=10, space_after=8)
    _ensure_style(doc, "CoverLetter.Signature", size=10, space_before=8)


def _set_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def _write_sample_assets(package: TemplatePackage) -> None:
    package_dir = package.directory
    sample_name = (
        "sample_resume.json" if package.document_type == "resume" else "sample_cover_letter.json"
    )
    sample_path = package_dir / sample_name
    if not sample_path.exists():
        sample_path.write_text("{}\n", encoding="utf-8", newline="\n")
    style_lock = package_dir / "style.lock.json"
    style_lock.write_text(
        json.dumps(
            {
                "template_id": package.template_id,
                "document_type": package.document_type,
                "styles": package.manifest.styles,
                "blocks": package.manifest.blocks,
            },
            indent=2,
        )
        + "\n",
        encoding="utf-8",
        newline="\n",
    )


def _ensure_required_markers(package: TemplatePackage) -> None:
    if not package.manifest.blocks:
        return
    if package.manifest.renderer == "latex":
        return
    try:
        doc = Document(str(package.template_path))
    except Exception:
        return
    text = _document_text(doc)
    changed = False
    if package.document_type == "resume":
        style = package.manifest.styles.get("normal")
    else:
        style = package.manifest.styles.get("body")
    for marker in package.manifest.blocks.values():
        if marker and marker not in text:
            _add_marker_paragraph(doc, marker, style)
            changed = True
    if changed:
        doc.save(str(package.template_path))


def _add_marker_paragraph(doc: Document, marker: str, style: str | None) -> None:
    if style:
        try:
            doc.add_paragraph(marker, style=style)
            return
        except KeyError:
            pass
    doc.add_paragraph(marker)


def _document_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def _slugify(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")[:60] or "template"


def _template_package_dir(document_type: str, template_id: str, template_root: Path) -> Path:
    if not _TEMPLATE_ID_RE.match(template_id):
        raise ValueError("Invalid template id.")

    type_root = (template_root / document_type).resolve()
    package_dir = (type_root / template_id).resolve()
    try:
        package_dir.relative_to(type_root)
    except ValueError as exc:
        raise ValueError("Invalid template id.") from exc
    return package_dir


def _public_asset_path(path: Path) -> str | None:
    if not path.exists():
        return None
    try:
        return path.resolve().relative_to(PROJECT_ROOT).as_posix()
    except ValueError:
        return path.name


def _unique_template_id(root: Path, template_id: str) -> str:
    candidate = template_id
    index = 2
    while (root / candidate).exists():
        candidate = f"{template_id}_{index}"
        index += 1
    return candidate
