"""Resume importer.

Parses existing Word (.docx) or PDF resumes into structured YAML format
using Claude CLI for intelligent extraction.
"""

from __future__ import annotations

import logging
from pathlib import Path

import yaml

from src.utils.llm import claude_generate

logger = logging.getLogger("autoapply.memory.resume_importer")

EXTRACTION_SYSTEM_PROMPT = """You are a resume parser. Extract structured data from the resume text provided.
Return ONLY valid YAML (no markdown fences, no explanations) matching this exact schema:

identity:
  full_name: "..."
  email: "..."
  phone: "..."
  location: "..."
  linkedin_url: "..."
  github_url: "..."
  portfolio_url: "..."

education:
  - institution: "..."
    degree: "..."
    field: "..."
    location: "..."
    start_date: "YYYY-MM"
    end_date: "YYYY-MM"
    gpa: "..."
    relevant_courses:
      - name: "..."
        tags: ["..."]

work_experiences:
  - company: "..."
    title: "..."
    location: "..."
    start_date: "YYYY-MM"
    end_date: "YYYY-MM"
    bullets:
      - text: "exact bullet text from resume"
        tags: ["skill1", "skill2"]

projects:
  - name: "..."
    role: "..."
    description: "..."
    tech_stack: ["..."]
    bullets:
      - text: "exact bullet text from resume"
        tags: ["skill1", "skill2"]

skills:
  languages: ["..."]
  frameworks: ["..."]
  databases: ["..."]
  tools: ["..."]
  domains: ["..."]

Rules:
- Preserve original bullet text exactly — do not rephrase or embellish
- Tags should be lowercase single words or short phrases (e.g., "python", "distributed_systems", "api_design")
- Dates in YYYY-MM format. Use "Present" for current positions
- If information is not in the resume, omit that field entirely
- For skills, categorize into the groups shown. If unsure of category, put in "tools"
"""


def extract_text_from_docx(path: Path) -> str:
    """Extract plain text from a .docx file."""
    from docx import Document

    doc = Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract from tables (common in resume templates)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    return "\n".join(paragraphs)


def extract_text_from_pdf(path: Path) -> str:
    """Extract plain text from a PDF file."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(path))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError("PyMuPDF (fitz) is required for PDF parsing. Install with: uv add pymupdf")


def import_resume(resume_path: Path, output_path: Path | None = None) -> dict:
    """Import a resume file and convert to structured YAML.

    Args:
        resume_path: Path to .docx or .pdf resume file.
        output_path: Optional path to save the generated YAML.

    Returns:
        Parsed profile data as a dict.
    """
    suffix = resume_path.suffix.lower()
    if suffix == ".docx":
        raw_text = extract_text_from_docx(resume_path)
    elif suffix == ".pdf":
        raw_text = extract_text_from_pdf(resume_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use .docx or .pdf")

    if not raw_text.strip():
        raise ValueError(f"No text extracted from {resume_path}")

    logger.info("Extracted %d chars from %s", len(raw_text), resume_path.name)

    # Use Claude CLI to parse into structured YAML
    prompt = f"Parse this resume into structured YAML:\n\n{raw_text}"
    response = claude_generate(prompt, system=EXTRACTION_SYSTEM_PROMPT, timeout=180)

    # Clean response — remove markdown fences if present
    cleaned = response.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split("\n")
        # Remove first and last fence lines
        lines = [l for l in lines if not l.strip().startswith("```")]
        cleaned = "\n".join(lines)

    # Parse YAML
    try:
        profile_data = yaml.safe_load(cleaned)
    except yaml.YAMLError as e:
        logger.error("Failed to parse LLM response as YAML: %s", e)
        logger.debug("Raw response:\n%s", cleaned)
        raise ValueError(f"LLM returned invalid YAML: {e}")

    if not isinstance(profile_data, dict):
        raise ValueError(f"Expected dict from YAML parse, got {type(profile_data)}")

    # Save to file if output path specified
    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            yaml.dump(profile_data, f, default_flow_style=False, allow_unicode=True, sort_keys=False)
        logger.info("Saved structured profile to %s", output_path)

    return profile_data
